from django.shortcuts import render
from django.http import HttpResponse

# Create your views here.

#importing the local preprocessing function
from .nepali_tfidf import calculate_tfidfsimilarity
from .predict import predict_lab
from .finger_sim import fingerprint_similarity
from .word_sim import compute_wordsimilarity
from .ngrams_sim import overlap_similarity
from .models import tfiles, thesis_docx
from .tfidf_sim import calculate_tfidfsimilarity

def home(request):
    tfilesData = tfiles.objects.all()
    for d in tfilesData:
        print(d.first_file.name)
    context = {
        "variable1" : "mishan thapa",
        "result" : "kshetri"
    }
    return render(request, 'home.html',context)

def plag_check(request):
    return render(request, 'plag_check.html')

def check(request):
    #lsa_sim = None
    if request.method == "POST":
        print("vayo")
        str1 = (request.POST["paragraph1"])
        str2 = (request.POST["paragraph2"])
        lsa_sim = calculate_tfidfsimilarity(str1,str2)
        #finger_sim = fingerprint_similarity(str1,str2)
        ngram_sim = overlap_similarity(str1,str2,2)
        label = predict_lab(lsa_sim,ngram_sim)
        sims = {
        "lsa_similarity" : lsa_sim,
        "ngram_sim" : ngram_sim,
        "plag_label" : label
        }
        return render(request, 'result.html', sims)



def tryy(request):
    return render(request, 'try.html')



def upload(request):
    if request.method == 'POST':
        print(request.FILES)
        f1 = request.FILES.get("fileone")
        #print(f1)
        f2 = request.FILES.get("filetwo")
        #data = tfiles(first_file = f1, second_file = f2)
        #data.save()
        f1_text = f1.read()
        f2_text = f2.read()
        f11_text = f1_text.decode()
        f22_text = f2_text.decode()
        #print(type(f11_text))
        #print(f11_text)
        ngram_sim = overlap_similarity(f11_text,f22_text)
        #lsa_sim = calculate_tfidfsimilarity(f11_text,f22_text)
        tfidf_sim = calculate_tfidfsimilarity(f11_text,f22_text)
        finger_sim = fingerprint_similarity(f11_text,f22_text)
        word_sim = compute_wordsimilarity(f11_text,f22_text)
        label = predict_lab(ngram_sim,finger_sim,word_sim,tfidf_sim)
        sims = {
        "tfidf_sim" : tfidf_sim,
        "ngram_sim" : ngram_sim,
        "finger_sim": finger_sim,
        "word_sim": word_sim,
        "plag_label" : label,
        }
        print("ata hai ata:")
        print(label)
        return render(request, 'result.html', sims)

    #return HttpResponse("success")
        
def upload_pdf(request):
    if request.method == 'POST' and request.FILES['pdf_file']:
        print("run vayo hai")
        uploaded_file = request.FILES['pdf_file']
        print(uploaded_file)
        return HttpResponse("File uploaded successfully.")
    else:
        return render(request, 'upload_pdf.html')
    
def thesis_upload_page(request):
    return render(request, 'thesis_upload.html')

def thesis_upload(request):
    if request.method == 'POST':
        print(request.FILES)
        got_thesis_file = request.FILES.get("thesis_file_docx")
        print(got_thesis_file)
        thesis_data = thesis_docx(thesis = got_thesis_file)
        thesis_data.save()
        return HttpResponse("hello.")
    

from docx import Document


def view_thesis(request):
    # thesis_docx_data = thesis_docx.objects.all()
    # for a in thesis_docx_data:
    #     print(a.thesis.name)
    # return HttpResponse("vayo")
    thesis_instance = thesis_docx.objects.get(thesis__icontains='short1.docx')
        
    docx_path = thesis_instance.thesis.path
    doc = Document(docx_path)

    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"
    print(content)
    return HttpResponse("vayo")


#to get the docx file from user and check their similarity


from django.shortcuts import render
from docx import Document
from .forms import DocxUploadForm
from .highlight_paragraph import process_plagiarized_paragraphs
import io
from .docx_to_paragraph import count_total_words,count_total_words_one_paragraph

from .new_highlight import highlight_new_wala

def extract_paragraphs(request):
    if request.method == 'POST':
        form = DocxUploadForm(request.POST, request.FILES)
        if form.is_valid():
            docx_file = request.FILES['docx_file']
            docx_file_title_name = docx_file.name
            docx_file_author_name = 'Inputed_Thesis_Author_Name'
            print(type(docx_file))
            sus_document = Document(docx_file)

            print(type(sus_document))
            input_paragraphs = [para.text.strip() for para in sus_document.paragraphs if para.text.strip()]  # Filter out empty or whitespace paragraphs
            # paragraphs = get_paragraphs_from_word_file(docx_file)
            print(input_paragraphs)
            # print(type(paragraphs))
            #return render(request, 'result.html', {'paragraphs': paragraphs})

            docx_files = thesis_docx.objects.all()
            file_names = [docx_file.thesis.name for docx_file in docx_files]
            print(file_names)
            print(file_names[0])
            paragraphs_list = []

            for docx_file in docx_files:
                document = Document(docx_file.thesis)
                #paragraphs = [para.text for para in document.paragraphs]
                para = [para.text.strip() for para in document.paragraphs if para.text.strip()]  # Filter out empty or whitespace paragraphs
                #paragraphs = get_paragraphs_from_word_file(document)
                paragraphs_list.append(para)


            plagiarised_paragraphs =[]
            # Loop through both lists simultaneously and call the function
            for para1 in input_paragraphs:
                paragraphs_list_counter =0
                for para2_group in paragraphs_list:
                    max_avg_feature_sim =0
                    max_sim_database_paragraph = ''
                    is_label_one = 0
                    for para2 in para2_group:
                        #asma two paragraph aaucha
                        ngram_sim = overlap_similarity(para1,para2)
                        tfidf_sim = calculate_tfidfsimilarity(para1,para2)
                        finger_sim = fingerprint_similarity(para1,para2)
                        word_sim = compute_wordsimilarity(para1,para2)
                        label = predict_lab(ngram_sim,finger_sim,word_sim,tfidf_sim)
                        avg_feature_sim = (ngram_sim + tfidf_sim + finger_sim + word_sim)/4
                        if(avg_feature_sim > max_avg_feature_sim and label == 1):
                            max_avg_feature_sim = avg_feature_sim
                            max_sim_database_paragraph = para2
                        
                        if(label == 1):
                            is_label_one = 1
                    if is_label_one ==1 :
                        my_dict = {"paragraph": para1 ,"database_paragraph":max_sim_database_paragraph, "source":file_names[paragraphs_list_counter], "average_feature_score": max_avg_feature_sim}
                        plagiarised_paragraphs.append(my_dict)
                        #asma lekhne sentence wala code
                        #
                    is_label_one = 0
                    paragraphs_list_counter = paragraphs_list_counter + 1
                #         print(para1)
                #         print("-----")
                #         print(para2)
                #         print("label: ")
                #         print(label)
                #     print("changeeeeeeeeeeee1111111")
                # print("changeeeeeeeeeeee22222222")
            print(plagiarised_paragraphs)
            #return render(request, 'try.html', {'paragraphs_list': paragraphs_list})
            #for highlighting docx
            #output_file_path = "highlighted_document.docx"  # Replace with the desired output file path
            #target_paragraphs = plagiarised_paragraphs
            #process_plagiarized_paragraphs(sus_document, output_file_path, target_paragraphs)
            final_plagiarised_paragraphs = {}
            for entry in plagiarised_paragraphs:
                paragraph = entry['paragraph']
                if paragraph in final_plagiarised_paragraphs:
                    if entry['average_feature_score'] > final_plagiarised_paragraphs[paragraph]['average_feature_score']:
                        final_plagiarised_paragraphs[paragraph] = entry
                else:
                    final_plagiarised_paragraphs[paragraph] = entry

            #for highlighting docx
            final_plagiarised_paragraphs = list(final_plagiarised_paragraphs.values())
            print(final_plagiarised_paragraphs)
            total_words_input_docx = count_total_words(input_paragraphs)
            print(total_words_input_docx)
            # Calculate and store the similarity index for each paragraph
            similarity_data = []
            for entry in final_plagiarised_paragraphs:
                sim_index = round((count_total_words_one_paragraph(entry['paragraph'])/ total_words_input_docx)*100,2)
                similarity_data.append({'input_paragraph':entry['paragraph'],'source': entry['source'],'database_paragraph':entry['database_paragraph'], 'sim_index': sim_index})
            print(similarity_data)

            #final wala dictionary aba
            grouped_paragraphs = []
            for entry in similarity_data:
                source = entry['source']
                input_paragraph = entry['input_paragraph']
                database_paragraph = entry['database_paragraph']
                sim_index = entry['sim_index']
                
                # Initialize a flag to False
                contains_source = False
                for single_dict in grouped_paragraphs:
                    if 'source' in single_dict and single_dict['source'] == source:
                        single_dict['input_paragraphs'].append(input_paragraph)
                        single_dict['database_paragraphs'].append(database_paragraph)
                        single_dict['percentage'] += sim_index
                        contains_source = True
                        break
                if(contains_source == False):
                    grouped_paragraphs.append({'input_paragraphs':[entry['input_paragraph']],'source': entry['source'],'database_paragraphs':[entry['database_paragraph']], 'percentage': sim_index, 'author': 'Author_name'})
            print(grouped_paragraphs)

            #highlight wala
            highlight_new_wala(grouped_paragraphs,sus_document,docx_file_title_name,docx_file_author_name)
            
            context = {'result_list': similarity_data}
            return render(request, 'try1.html', context)
    else:
        form = DocxUploadForm()
    return render(request, 'upload.html', {'form': form})

