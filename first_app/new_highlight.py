import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
# Load an existing DOCX document
#doc = docx.Document('short2.docx')



def random_rgb_color():
    while True:
        red = random.randint(150, 255)  # Adjusted the range for higher brightness
        green = random.randint(100, 255)  # Adjusted the range for higher brightness
        blue = random.randint(100, 255) 
        return red,green,blue
    

def find_common_sequences(string1, string2,threshold):
    # Split the strings into words
    words1 = string1.split(' ')
    words2 = string2.split()

    # Initialize variables to keep track of common sequences
    common_sequences = []
    current_sequence = []

    # Iterate through words in the first string
    checked_word = ''
    
    for word1 in words1:
        #print(checked_word)
        if word1!= '':
        # If the word is in the second string
            if word1 in words2:
                current_sequence.append(word1)
                if len(current_sequence) == 1:
                    current_start_index = len(checked_word)
            else:
                # Check if the current sequence is longer than the threshold
                if len(current_sequence) >= threshold:
                    text = ' '.join(current_sequence)
                    print(text)
                    print(len(text))
                    common_sequences.append({
                        'sequence': ' '.join(current_sequence),
                        'start_index': current_start_index,
                        'end_index': len(checked_word) - 1
                    })
                current_sequence = []
        checked_word = checked_word+ ' ' + word1
        
    # Check if the current sequence at the end is longer than the threshold
    if len(current_sequence) >= threshold:
        common_sequences.append({
            'sequence': ' '.join(current_sequence),
            'start_index': current_start_index,
            'end_index': len(checked_word) - 1
        })
        
    result_list = []
    print(common_sequences)
    # Initialize a variable to keep track of the end value of the previous dictionary
    prev_end = None
    if common_sequences:
        if common_sequences[0]['start_index'] != 0:
            result_list.append({'copied': False,'sen_pos':[0, common_sequences[0]['start_index']]})
        # Iterate through the dictionaries and extract 'start' and 'end' values
        for seq_info in common_sequences:
            start = seq_info['start_index']
            end = seq_info['end_index']

            # Include the missing range
            if prev_end is not None and start > prev_end:
                result_list.append({'copied': False,'sen_pos':[prev_end, start]})

            result_list.append({'copied': True,'sen_pos':[start, end]})
            prev_end = end

        return result_list
    else:
        return result_list
    

def highlight_new_wala(final_plagiarised_paragraphs_grouped_sources,doc,docx_file_title_name,docx_file_author_name):
    # Define the text you want to change the color of
    color = []
    number_of_source = 0
    total_plagiarism = 0
    for sources in final_plagiarised_paragraphs_grouped_sources:
        number_of_source = number_of_source +1
        red,green,blue = random_rgb_color()
        color.append((red,green,blue))
        total_plagiarism = total_plagiarism + sources['percentage']
        count = 0
        for target_paragraph in sources['input_paragraphs']:
            input_para = target_paragraph # yo rw tala ko paragraph use garni sentence level index patta lauda
            database_para = sources['database_paragraphs'][count]
            count = count + 1
            print(target_paragraph)
            print("\n")
            
            # Iterate through paragraphs to find the paragraph containing the specific text
            for paragraph in doc.paragraphs:
                print(paragraph.text)
                print("\n")
                # Search for the target text within the paragraph
                if target_paragraph in paragraph.text:
                    para_text = paragraph.text
                    print(paragraph.text)
                    print("\n")
                    # Clear the original paragraph text
                    paragraph.clear()
                    #call sentence pattalauni function
                    indexes = find_common_sequences(input_para, database_para, threshold=2)
                    for index in indexes:
                        if index['copied'] is True:
                            print(index['sen_pos'])
                            run = paragraph.add_run(para_text[index['sen_pos'][0]:index['sen_pos'][1]])
                            run.font.color.rgb = RGBColor(red, green, blue)
                            run.font.underline = WD_UNDERLINE.THICK
                        else:
                            run = paragraph.add_run(para_text[index['sen_pos'][0]:index['sen_pos'][1]])
                            run.font.color.rgb = RGBColor(red, green, blue)
                    if indexes[-1]['sen_pos'][1]!= len(para_text):
                        run = paragraph.add_run(para_text[index['sen_pos'][1]:len(para_text)])
                        run.font.color.rgb = RGBColor(red, green, blue)

                    break

    doc.add_page_break()

    blue_color = RGBColor(0, 0, 255) 
    # Add a Title to the document
    heading = doc.add_paragraph('Plagiarism Report')
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set alignment
    heading.runs[0].bold = True  # Make the first run (the entire paragraph) bold
    heading.runs[0].font.size = Pt(30) 
    heading.runs[0].font.color.rgb = blue_color
    # Add Thesis Details heading
    #doc.add_heading('Thesis Details').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    report = doc.add_paragraph('Thesis Details')
    report.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set alignment
    report.runs[0].bold = True  # Make the first run (the entire paragraph) bold
    report.runs[0].font.size = Pt(22) 
    report.runs[0].font.color.rgb = blue_color


    # Create a table with 2 rows and 2 columns
    table = doc.add_table(rows=2, cols=2)

    # Set the alignment of the table to center
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add content to the table cells
    table.cell(0, 0).text = "Thesis Title"
    table.cell(0,0).paragraphs[0].runs[0].bold = True
    table.cell(0,0).paragraphs[0].runs[0].font.size = Pt(14) 
    table.cell(0, 1).text = docx_file_title_name
    table.cell(1, 0).text = "Author"
    table.cell(1,0).paragraphs[0].runs[0].bold = True
    table.cell(1,0).paragraphs[0].runs[0].font.size = Pt(14) 
    table.cell(1, 1).text = docx_file_author_name

    # Add Plagiarism Analysis heading
    #doc.add_heading('Plagiarism Analysis').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    analysis = doc.add_paragraph('Plagiarism Analysis')
    analysis.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set alignment
    analysis.runs[0].bold = True  # Make the first run (the entire paragraph) bold
    analysis.runs[0].font.size = Pt(22) 
    analysis.runs[0].font.color.rgb = blue_color
    # to display total plag percentage
    ana_table = doc.add_table(rows=1, cols=2)
    ana_table.cell(0,0).text = "Plagiarized Percentage"
    ana_table.cell(0,0).paragraphs[0].runs[0].bold = True
    ana_table.cell(0,0).paragraphs[0].runs[0].font.size = Pt(14) 
    ana_table.cell(0,1).text = str(total_plagiarism) +"%"

    #doc.add_heading('Plagiarism Analysis').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sources = doc.add_paragraph('Sources')
    sources.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set alignment
    sources.runs[0].bold = True  # Make the first run (the entire paragraph) bold
    sources.runs[0].font.size = Pt(22) 
    sources.runs[0].font.color.rgb = blue_color
    # table to display sources
    source_table = doc.add_table(rows=number_of_source+1, cols=4)
    print(number_of_source)
    source_table.cell(0,0).text = "Index"
    source_table.cell(0,1).text = "Source Thesis"
    source_table.cell(0,2).text = "Author"
    source_table.cell(0,3).text = "Plagiarized percentage"
    # Define the table style with borders
    #table.style = 'Light Shading'
    for cell in source_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(14)
    # Add content to the cells (optional)
    for i in range(1,number_of_source+1):
        source_table.cell(i,0).text = str(i)
        source_table.cell(i,1).text = final_plagiarised_paragraphs_grouped_sources[i-1]['source']
        source_table.cell(i,2).text = final_plagiarised_paragraphs_grouped_sources[i-1]['author']+str(i)
        source_table.cell(i,3).text = str(final_plagiarised_paragraphs_grouped_sources[i-1]['percentage']) + '%'
        
        for cell in source_table.rows[i].cells:
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(color[i-1][0], color[i-1][1], color[i-1][2])
    doc.save('short2_modified.docx')
